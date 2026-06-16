"""
Text extraction module for PDF and DOCX resume files.
Uses pdfplumber for PDFs (with PyPDF2 fallback) and python-docx for DOCX files.
"""

import os
import re
from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_file(filepath: str) -> str:
    """Extract text from a PDF or DOCX file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber, with PyPDF2 fallback."""
    text = ""

    # Try pdfplumber first (better quality)
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if text.strip():
            logger.info(f"PDF text extracted with pdfplumber: {len(text)} chars")
            return clean_text(text)
    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")

    # Fallback to PyPDF2
    try:
        import PyPDF2
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if text.strip():
            logger.info(f"PDF text extracted with PyPDF2: {len(text)} chars")
            return clean_text(text)
    except Exception as e:
        logger.error(f"PyPDF2 also failed: {str(e)}")

    if not text.strip():
        raise ValueError("Could not extract text from PDF. The file may be scanned or corrupted.")

    return clean_text(text)


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        text = ""

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        if not text.strip():
            raise ValueError("DOCX file appears to be empty")

        logger.info(f"DOCX text extracted: {len(text)} chars")
        return clean_text(text)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove excessive spaces
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable characters (keep newlines and tabs)
    text = re.sub(r'[^\S\n\t]+', ' ', text)
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()
